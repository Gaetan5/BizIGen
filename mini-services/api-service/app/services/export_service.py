"""
BizGen AI - Export Service
Handles PDF, DOCX, PNG generation and exports with professional styling
"""
import json
import io
import os
import zipfile
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime
from pathlib import Path
import math

# PDF Generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
    PageBreak, Image, KeepTogether, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.graphics.shapes import Drawing, Rect, String
from reportlab.graphics import renderPDF

# DOCX Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Image Generation
from PIL import Image as PILImage, ImageDraw, ImageFont, ImageColor

from app.config import settings


# ============================================
# CONSTANTS & COLORS
# ============================================

# Brand Colors
BIZGEN_PRIMARY = "#1E40AF"      # Dark Blue
BIZGEN_SECONDARY = "#3B82F6"    # Blue
BIZGEN_ACCENT = "#10B981"       # Green
BIZGEN_DARK = "#1F2937"         # Dark Gray
BIZGEN_LIGHT = "#F3F4F6"        # Light Gray
BIZGEN_WHITE = "#FFFFFF"

# Canvas Block Colors for PNG
BMC_COLORS = {
    "key_partners": "#E0F2FE",          # Light Blue
    "key_activities": "#DBEAFE",        # Blue Tint
    "key_resources": "#BFDBFE",         # Blue Light
    "value_propositions": "#FEF3C7",    # Yellow/Amber
    "customer_relationships": "#FCE7F3", # Pink
    "channels": "#F3E8FF",              # Purple
    "customer_segments": "#D1FAE5",     # Green
    "cost_structure": "#FEE2E2",        # Red
    "revenue_streams": "#ECFDF5",       # Green Light
}

LEAN_COLORS = {
    "problem": "#FEE2E2",               # Red
    "existing_alternatives": "#FECACA", # Red Light
    "solution": "#D1FAE5",              # Green
    "key_metrics": "#DBEAFE",           # Blue
    "unique_value_proposition": "#FEF3C7",  # Yellow
    "high_level_concept": "#FDE68A",    # Yellow Dark
    "unfair_advantage": "#E0E7FF",      # Indigo
    "channels": "#F3E8FF",              # Purple
    "customer_segments": "#FCE7F3",     # Pink
    "cost_structure": "#FED7AA",        # Orange
    "revenue_streams": "#ECFDF5",       # Green Light
}

# Default fonts (will use system fonts)
FONT_TITLE = "Helvetica-Bold"
FONT_HEADING = "Helvetica-Bold"
FONT_BODY = "Helvetica"


class ExportService:
    """Service for generating exports in various formats"""
    
    def __init__(self):
        self.export_dir = Path("./exports")
        self.export_dir.mkdir(exist_ok=True)
        
        # Load fonts for PIL
        self._init_fonts()
    
    def _init_fonts(self):
        """Initialize fonts for PIL image generation"""
        self.pil_fonts = {
            'title': None,
            'heading': None,
            'body': None,
            'small': None
        }
        
        # Try to load fonts (fallback to default)
        font_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/System/Library/Fonts/Helvetica.ttc",
            "/System/Library/Fonts/SFNS.ttf",
            "C:\\Windows\\Fonts\\arial.ttf",
        ]
        
        for path in font_paths:
            if os.path.exists(path):
                try:
                    self.pil_fonts['title'] = ImageFont.truetype(path, 24)
                    self.pil_fonts['heading'] = ImageFont.truetype(path, 14)
                    self.pil_fonts['body'] = ImageFont.truetype(path, 11)
                    self.pil_fonts['small'] = ImageFont.truetype(path, 9)
                    break
                except Exception:
                    continue
        
        # Fallback to default font
        if self.pil_fonts['title'] is None:
            self.pil_fonts['title'] = ImageFont.load_default()
            self.pil_fonts['heading'] = ImageFont.load_default()
            self.pil_fonts['body'] = ImageFont.load_default()
            self.pil_fonts['small'] = ImageFont.load_default()

    # ============================================
    # PNG GENERATION - CANVAS
    # ============================================
    
    def generate_bmc_png(self, bmc_data: Dict[str, Any], project_name: str) -> bytes:
        """
        Generate Business Model Canvas as PNG image
        
        Layout:
        ┌─────────────┬─────────────┬─────────────┬─────────────┐
        │ Key         │ Key         │ Value       │ Customer   │
        │ Partners    │ Activities  │ Prop.       │ Relations  │
        ├─────────────┼─────────────┤             ├─────────────┤
        │ Key         │             │             │ Channels   │
        │ Resources   │             │             │            │
        ├─────────────┴─────────────┴─────────────┴─────────────┤
        │                    Customer Segments                  │
        ├───────────────────────────┬───────────────────────────┤
        │       Cost Structure      │     Revenue Streams       │
        └───────────────────────────┴───────────────────────────┘
        """
        # Canvas dimensions
        width = 1400
        height = 900
        margin = 20
        header_height = 60
        footer_height = 30
        
        # Create image
        img = PILImage.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        
        # Draw header
        draw.rectangle([0, 0, width, header_height], fill=BIZGEN_PRIMARY)
        draw.text(
            (width // 2, header_height // 2),
            f"Business Model Canvas - {project_name}",
            font=self.pil_fonts['title'],
            fill='white',
            anchor='mm'
        )
        
        # Calculate grid dimensions
        grid_top = header_height + margin
        grid_bottom = height - footer_height - margin
        grid_left = margin
        grid_right = width - margin
        grid_width = grid_right - grid_left
        grid_height = grid_bottom - grid_top
        
        # Row heights
        row1_height = int(grid_height * 0.35)
        row2_height = int(grid_height * 0.35)
        row3_height = int(grid_height * 0.15)
        row4_height = grid_height - row1_height - row2_height - row3_height
        
        # Column widths (proportional to BMC standard layout)
        col1_width = int(grid_width * 0.20)  # Key Partners
        col2_width = int(grid_width * 0.20)  # Key Activities / Resources
        col3_width = int(grid_width * 0.25)  # Value Propositions (center)
        col4_width = int(grid_width * 0.20)  # Customer Relations / Channels
        col5_width = grid_width - col1_width - col2_width - col3_width - col4_width
        
        # Current positions
        y = grid_top
        x = grid_left
        
        # Block labels and data mapping
        bmc_blocks = [
            # Row 1
            {
                "key": "key_partners",
                "title": "Partenaires Clés",
                "rect": (x, y, x + col1_width, y + row1_height + row2_height),
                "color": BMC_COLORS["key_partners"]
            },
            {
                "key": "key_activities",
                "title": "Activités Clés",
                "rect": (x + col1_width, y, x + col1_width + col2_width, y + row1_height),
                "color": BMC_COLORS["key_activities"]
            },
            {
                "key": "value_propositions",
                "title": "Propositions de Valeur",
                "rect": (x + col1_width + col2_width, y, x + col1_width + col2_width + col3_width, y + row1_height + row2_height),
                "color": BMC_COLORS["value_propositions"]
            },
            {
                "key": "customer_relationships",
                "title": "Relations Clients",
                "rect": (x + col1_width + col2_width + col3_width, y, x + col1_width + col2_width + col3_width + col4_width, y + row1_height),
                "color": BMC_COLORS["customer_relationships"]
            },
            {
                "key": "channels",
                "title": "Canaux",
                "rect": (x + col1_width + col2_width + col3_width, y + row1_height, x + col1_width + col2_width + col3_width + col4_width + col5_width, y + row1_height + row2_height),
                "color": BMC_COLORS["channels"]
            },
            # Row 2
            {
                "key": "key_resources",
                "title": "Ressources Clés",
                "rect": (x + col1_width, y + row1_height, x + col1_width + col2_width, y + row1_height + row2_height),
                "color": BMC_COLORS["key_resources"]
            },
            # Row 3 - Customer Segments (full width minus Cost/Revenue columns)
            {
                "key": "customer_segments",
                "title": "Segments Clients",
                "rect": (x, y + row1_height + row2_height, x + grid_width, y + row1_height + row2_height + row3_height),
                "color": BMC_COLORS["customer_segments"]
            },
            # Row 4 - Cost Structure & Revenue Streams
            {
                "key": "cost_structure",
                "title": "Structure des Coûts",
                "rect": (x, y + row1_height + row2_height + row3_height, x + grid_width // 2, y + grid_height),
                "color": BMC_COLORS["cost_structure"]
            },
            {
                "key": "revenue_streams",
                "title": "Sources de Revenus",
                "rect": (x + grid_width // 2, y + row1_height + row2_height + row3_height, x + grid_width, y + grid_height),
                "color": BMC_COLORS["revenue_streams"]
            },
        ]
        
        # Draw each block
        for block in bmc_blocks:
            self._draw_canvas_block(draw, block, bmc_data.get(block["key"], []))
        
        # Draw footer
        draw.rectangle([0, height - footer_height, width, height], fill=BIZGEN_LIGHT)
        draw.text(
            (width // 2, height - footer_height // 2),
            f"Généré par BizGen AI - {datetime.now().strftime('%d/%m/%Y')}",
            font=self.pil_fonts['small'],
            fill=BIZGEN_DARK,
            anchor='mm'
        )
        
        # Convert to bytes
        buffer = io.BytesIO()
        img.save(buffer, format='PNG', dpi=(150, 150))
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_lean_canvas_png(self, lean_data: Dict[str, Any], project_name: str) -> bytes:
        """
        Generate Lean Canvas as PNG image
        
        Layout:
        ┌───────────────┬───────────────┬───────────────┬───────────────┬───────────────┐
        │   Problem     │   Solution    │   UVP         │ Unfair        │ Customer      │
        │               │               │               │ Advantage     │ Segments      │
        ├───────────────┼───────────────┤   High-Level  ├───────────────┼───────────────┤
        │ Existing      │ Key Metrics   │   Concept     │               │               │
        │ Alternatives  │               │               │               │               │
        ├───────────────┴───────────────┴───────────────┴───────────────┴───────────────┤
        │                              Channels                                          │
        ├───────────────────────────────────┬───────────────────────────────────────────┤
        │         Cost Structure            │          Revenue Streams                  │
        └───────────────────────────────────┴───────────────────────────────────────────┘
        """
        # Canvas dimensions
        width = 1400
        height = 900
        margin = 20
        header_height = 60
        footer_height = 30
        
        # Create image
        img = PILImage.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        
        # Draw header
        draw.rectangle([0, 0, width, header_height], fill='#059669')  # Green for Lean Canvas
        draw.text(
            (width // 2, header_height // 2),
            f"Lean Canvas - {project_name}",
            font=self.pil_fonts['title'],
            fill='white',
            anchor='mm'
        )
        
        # Calculate grid dimensions
        grid_top = header_height + margin
        grid_bottom = height - footer_height - margin
        grid_left = margin
        grid_right = width - margin
        grid_width = grid_right - grid_left
        grid_height = grid_bottom - grid_top
        
        # Row heights
        row1_height = int(grid_height * 0.40)
        row2_height = int(grid_height * 0.25)
        row3_height = int(grid_height * 0.15)
        row4_height = grid_height - row1_height - row2_height - row3_height
        
        # Column widths (5 columns)
        col_width = grid_width // 5
        
        # Current positions
        y = grid_top
        x = grid_left
        
        # Lean Canvas blocks
        lean_blocks = [
            # Row 1
            {
                "key": "problem",
                "title": "Problème",
                "rect": (x, y, x + col_width, y + row1_height + row2_height),
                "color": LEAN_COLORS["problem"]
            },
            {
                "key": "solution",
                "title": "Solution",
                "rect": (x + col_width, y, x + 2 * col_width, y + row1_height),
                "color": LEAN_COLORS["solution"]
            },
            {
                "key": "unique_value_proposition",
                "title": "Proposition de Valeur Unique",
                "rect": (x + 2 * col_width, y, x + 3 * col_width, y + row1_height),
                "color": LEAN_COLORS["unique_value_proposition"]
            },
            {
                "key": "unfair_advantage",
                "title": "Avantage Déloyal",
                "rect": (x + 3 * col_width, y, x + 4 * col_width, y + row1_height + row2_height),
                "color": LEAN_COLORS["unfair_advantage"]
            },
            {
                "key": "customer_segments",
                "title": "Segments Clients",
                "rect": (x + 4 * col_width, y, x + 5 * col_width, y + row1_height + row2_height),
                "color": LEAN_COLORS["customer_segments"]
            },
            # Row 2
            {
                "key": "existing_alternatives",
                "title": "Alternatives Existantes",
                "rect": (x + col_width, y + row1_height, x + 2 * col_width, y + row1_height + row2_height),
                "color": LEAN_COLORS["existing_alternatives"]
            },
            {
                "key": "high_level_concept",
                "title": "Concept de Haut Niveau",
                "rect": (x + 2 * col_width, y + row1_height, x + 3 * col_width, y + row1_height + row2_height),
                "color": LEAN_COLORS["high_level_concept"]
            },
            # Row 3 - Key Metrics (spans 2 columns)
            {
                "key": "key_metrics",
                "title": "Indicateurs Clés",
                "rect": (x, y + row1_height + row2_height, x + 2 * col_width, y + row1_height + row2_height + row3_height),
                "color": LEAN_COLORS["key_metrics"]
            },
            # Row 3 - Channels (spans 3 columns)
            {
                "key": "channels",
                "title": "Canaux",
                "rect": (x + 2 * col_width, y + row1_height + row2_height, x + 5 * col_width, y + row1_height + row2_height + row3_height),
                "color": LEAN_COLORS["channels"]
            },
            # Row 4 - Cost Structure & Revenue Streams
            {
                "key": "cost_structure",
                "title": "Structure des Coûts",
                "rect": (x, y + row1_height + row2_height + row3_height, x + grid_width // 2, y + grid_height),
                "color": LEAN_COLORS["cost_structure"]
            },
            {
                "key": "revenue_streams",
                "title": "Sources de Revenus",
                "rect": (x + grid_width // 2, y + row1_height + row2_height + row3_height, x + grid_width, y + grid_height),
                "color": LEAN_COLORS["revenue_streams"]
            },
        ]
        
        # Draw each block
        for block in lean_blocks:
            self._draw_canvas_block(draw, block, lean_data.get(block["key"], []))
        
        # Draw footer
        draw.rectangle([0, height - footer_height, width, height], fill=BIZGEN_LIGHT)
        draw.text(
            (width // 2, height - footer_height // 2),
            f"Généré par BizGen AI - {datetime.now().strftime('%d/%m/%Y')}",
            font=self.pil_fonts['small'],
            fill=BIZGEN_DARK,
            anchor='mm'
        )
        
        # Convert to bytes
        buffer = io.BytesIO()
        img.save(buffer, format='PNG', dpi=(150, 150))
        buffer.seek(0)
        return buffer.getvalue()
    
    def _draw_canvas_block(self, draw: ImageDraw.ImageDraw, block: Dict, data: Any):
        """Draw a single canvas block with title and content"""
        rect = block["rect"]
        color = block["color"]
        title = block["title"]
        
        # Draw block background
        draw.rectangle(rect, fill=color, outline=BIZGEN_DARK, width=1)
        
        # Draw title bar
        title_height = 25
        draw.rectangle(
            [rect[0], rect[1], rect[2], rect[1] + title_height],
            fill=BIZGEN_DARK,
            outline=BIZGEN_DARK
        )
        
        # Draw title text
        draw.text(
            (rect[0] + 5, rect[1] + title_height // 2),
            title,
            font=self.pil_fonts['heading'],
            fill='white',
            anchor='lm'
        )
        
        # Draw content
        content_y = rect[1] + title_height + 10
        content_x = rect[0] + 10
        max_width = rect[2] - rect[0] - 20
        
        # Format content based on data type
        if isinstance(data, list):
            items = data
        elif isinstance(data, dict):
            items = [f"{k}: {v}" for k, v in data.items()]
        elif isinstance(data, str):
            items = [data] if data else []
        else:
            items = []
        
        # Draw each item as bullet point
        line_height = 16
        for item in items[:10]:  # Limit to 10 items
            if content_y + line_height > rect[3] - 5:
                break
            
            item_text = str(item)
            if len(item_text) > 50:
                item_text = item_text[:47] + "..."
            
            draw.text(
                (content_x, content_y),
                f"• {item_text}",
                font=self.pil_fonts['body'],
                fill=BIZGEN_DARK
            )
            content_y += line_height

    # ============================================
    # PDF GENERATION - IMPROVED
    # ============================================
    
    def _add_watermark(self, canvas_obj, doc):
        """Add watermark for Free plan users"""
        canvas_obj.saveState()
        canvas_obj.setFont('Helvetica', 40)
        canvas_obj.setFillColor(colors.Color(0.8, 0.8, 0.8, alpha=0.3))
        canvas_obj.rotate(45)
        canvas_obj.drawString(200, 100, "BIZGEN AI - FREE VERSION")
        canvas_obj.restoreState()
    
    def _add_header_footer(self, canvas_obj, doc, project_name: str, is_free: bool = False):
        """Add header with logo and footer to PDF pages"""
        canvas_obj.saveState()
        
        # Header
        header_height = 40
        canvas_obj.setFillColor(colors.HexColor(BIZGEN_PRIMARY))
        canvas_obj.rect(0, doc.height + doc.topMargin, doc.width + doc.leftMargin + doc.rightMargin, header_height, fill=1, stroke=0)
        
        # Header text
        canvas_obj.setFillColor(colors.white)
        canvas_obj.setFont('Helvetica-Bold', 12)
        canvas_obj.drawString(doc.leftMargin, doc.height + doc.topMargin + 15, f"BizGen AI - {project_name}")
        
        # Page number in header
        canvas_obj.drawRightString(doc.width + doc.leftMargin, doc.height + doc.topMargin + 15, f"Page {doc.page}")
        
        # Footer
        canvas_obj.setFillColor(colors.HexColor(BIZGEN_DARK))
        canvas_obj.setFont('Helvetica', 8)
        canvas_obj.drawString(
            doc.leftMargin, 
            20, 
            f"Document généré par BizGen AI le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        )
        
        # Free plan watermark
        if is_free:
            canvas_obj.setFont('Helvetica-Bold', 8)
            canvas_obj.setFillColor(colors.HexColor('#EF4444'))
            canvas_obj.drawRightString(doc.width + doc.leftMargin, 20, "VERSION GRATUITE")
        
        canvas_obj.restoreState()
    
    def generate_bmc_pdf(
        self, 
        bmc_data: Dict[str, Any], 
        project_name: str,
        is_free_plan: bool = False
    ) -> bytes:
        """Generate Business Model Canvas as PDF with improved layout"""
        buffer = io.BytesIO()
        
        # Use landscape for better BMC layout
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            rightMargin=1.5*cm,
            leftMargin=1.5*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'BMCTitle',
            parent=styles['Heading1'],
            fontSize=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor(BIZGEN_PRIMARY),
            spaceAfter=20,
            fontName=FONT_TITLE
        )
        
        block_title_style = ParagraphStyle(
            'BlockTitle',
            parent=styles['Heading2'],
            fontSize=11,
            textColor=colors.white,
            alignment=TA_CENTER,
            fontName=FONT_HEADING
        )
        
        body_style = ParagraphStyle(
            'BMCBody',
            parent=styles['Normal'],
            fontSize=9,
            fontName=FONT_BODY,
            leading=12
        )
        
        story = []
        
        # Title with free plan indicator
        title_text = f"Business Model Canvas - {project_name}"
        if is_free_plan:
            title_text += " (Version Gratuite)"
        story.append(Paragraph(title_text, title_style))
        story.append(Spacer(1, 0.3*cm))
        
        # BMC Table Layout
        # Row 1: Key Partners, Key Activities, Value Props, Customer Relations, Channels placeholder
        # Row 2: Key Partners cont, Key Resources, Value Props cont, Channels, Customer Segments placeholder
        
        def create_block(title: str, items: Any, bg_color: str) -> List:
            """Create a styled table cell for a BMC block"""
            if isinstance(items, list):
                content = "<br/>".join([f"• {item}" for item in items[:8]]) if items else "-"
            elif isinstance(items, dict):
                content = "<br/>".join([f"• {k}: {v}" for k, v in list(items.items())[:8]]) if items else "-"
            else:
                content = str(items) if items else "-"
            
            inner_table = Table(
                [
                    [Paragraph(f"<b>{title}</b>", block_title_style)],
                    [Paragraph(content, body_style)]
                ],
                colWidths=[None]
            )
            inner_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BIZGEN_DARK)),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor(bg_color)),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor(BIZGEN_DARK)),
            ]))
            return inner_table
        
        # Create BMC blocks
        key_partners = create_block("Partenaires Clés", bmc_data.get("key_partners", []), BMC_COLORS["key_partners"])
        key_activities = create_block("Activités Clés", bmc_data.get("key_activities", []), BMC_COLORS["key_activities"])
        key_resources = create_block("Ressources Clés", bmc_data.get("key_resources", []), BMC_COLORS["key_resources"])
        value_props = create_block("Propositions de Valeur", bmc_data.get("value_propositions", []), BMC_COLORS["value_propositions"])
        cust_relations = create_block("Relations Clients", bmc_data.get("customer_relationships", []), BMC_COLORS["customer_relationships"])
        channels = create_block("Canaux", bmc_data.get("channels", []), BMC_COLORS["channels"])
        cust_segments = create_block("Segments Clients", bmc_data.get("customer_segments", []), BMC_COLORS["customer_segments"])
        
        # Cost structure
        cost_data = bmc_data.get("cost_structure", {})
        cost_items = []
        if isinstance(cost_data, dict):
            if "fixed_costs" in cost_data:
                for cost in cost_data["fixed_costs"]:
                    cost_items.append(f"{cost.get('item', '')}: {cost.get('amount', '')}")
            if "variable_costs" in cost_data:
                for cost in cost_data["variable_costs"]:
                    cost_items.append(f"{cost.get('item', '')}: {cost.get('amount', '')}")
        cost_structure = create_block("Structure des Coûts", cost_items, BMC_COLORS["cost_structure"])
        
        # Revenue streams
        revenue_data = bmc_data.get("revenue_streams", [])
        revenue_items = []
        for stream in revenue_data:
            if isinstance(stream, dict):
                revenue_items.append(f"{stream.get('source', '')}: {stream.get('pricing', '')}")
            else:
                revenue_items.append(str(stream))
        revenue_streams = create_block("Sources de Revenus", revenue_items, BMC_COLORS["revenue_streams"])
        
        # Main BMC table (5 columns layout)
        # Calculate widths
        page_width = landscape(A4)[0] - 3*cm
        col_width = page_width / 5
        
        bmc_table = Table([
            # Row 1
            [
                Table([[key_partners]], colWidths=[col_width]),
                Table([[key_activities]], colWidths=[col_width]),
                Table([[value_props]], colWidths=[col_width]),
                Table([[cust_relations]], colWidths=[col_width]),
                Table([[channels]], colWidths=[col_width]),
            ],
            # Row 2
            [
                Table([[""]], colWidths=[col_width]),  # Empty continuation
                Table([[key_resources]], colWidths=[col_width]),
                Table([[""]], colWidths=[col_width]),  # Empty continuation
                Table([[""]], colWidths=[col_width]),  # Empty continuation
                Table([[""]], colWidths=[col_width]),  # Empty continuation
            ],
            # Row 3 - Customer Segments (full width)
            [
                Table([[cust_segments]], colWidths=[page_width]),
            ],
            # Row 4 - Cost & Revenue
            [
                Table([[cost_structure]], colWidths=[page_width/2]),
                Table([[revenue_streams]], colWidths=[page_width/2]),
            ],
        ], colWidths=[col_width] * 5, rowHeights=[None, None, None, None])
        
        bmc_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(BIZGEN_DARK)),
        ]))
        
        story.append(bmc_table)
        
        # Footer
        story.append(Spacer(1, 0.5*cm))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph(f"Généré par BizGen AI - {datetime.now().strftime('%d/%m/%Y')}", footer_style))
        
        # Build PDF with header/footer
        def on_page(canvas_obj, doc):
            self._add_header_footer(canvas_obj, doc, project_name, is_free_plan)
            if is_free_plan:
                self._add_watermark(canvas_obj, doc)
        
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_lean_canvas_pdf(
        self,
        lean_data: Dict[str, Any],
        project_name: str,
        is_free_plan: bool = False
    ) -> bytes:
        """Generate Lean Canvas as PDF"""
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            rightMargin=1.5*cm,
            leftMargin=1.5*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'LeanTitle',
            parent=styles['Heading1'],
            fontSize=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#059669'),  # Green for Lean Canvas
            spaceAfter=20,
            fontName=FONT_TITLE
        )
        
        block_title_style = ParagraphStyle(
            'BlockTitle',
            parent=styles['Heading2'],
            fontSize=10,
            textColor=colors.white,
            alignment=TA_CENTER,
            fontName=FONT_HEADING
        )
        
        body_style = ParagraphStyle(
            'LeanBody',
            parent=styles['Normal'],
            fontSize=9,
            fontName=FONT_BODY,
            leading=12
        )
        
        story = []
        
        # Title
        title_text = f"Lean Canvas - {project_name}"
        if is_free_plan:
            title_text += " (Version Gratuite)"
        story.append(Paragraph(title_text, title_style))
        story.append(Spacer(1, 0.3*cm))
        
        # Create Lean Canvas blocks
        def create_block(title: str, items: Any, bg_color: str) -> Table:
            if isinstance(items, list):
                content = "<br/>".join([f"• {item}" for item in items[:8]]) if items else "-"
            elif isinstance(items, dict):
                content = "<br/>".join([f"• {k}: {v}" for k, v in list(items.items())[:8]]) if items else "-"
            elif isinstance(items, str):
                content = items if items else "-"
            else:
                content = str(items) if items else "-"
            
            inner_table = Table(
                [
                    [Paragraph(f"<b>{title}</b>", block_title_style)],
                    [Paragraph(content, body_style)]
                ]
            )
            inner_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BIZGEN_DARK)),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor(bg_color)),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor(BIZGEN_DARK)),
            ]))
            return inner_table
        
        page_width = landscape(A4)[0] - 3*cm
        col_width = page_width / 5
        
        # Create blocks
        problem = create_block("Problème", lean_data.get("problem", []), LEAN_COLORS["problem"])
        existing_alt = create_block("Alternatives Existantes", lean_data.get("existing_alternatives", []), LEAN_COLORS["existing_alternatives"])
        solution = create_block("Solution", lean_data.get("solution", []), LEAN_COLORS["solution"])
        key_metrics = create_block("Indicateurs Clés", lean_data.get("key_metrics", []), LEAN_COLORS["key_metrics"])
        uvp = create_block("Proposition de Valeur Unique", lean_data.get("unique_value_proposition", ""), LEAN_COLORS["unique_value_proposition"])
        high_level = create_block("Concept de Haut Niveau", lean_data.get("high_level_concept", ""), LEAN_COLORS["high_level_concept"])
        unfair_adv = create_block("Avantage Déloyal", lean_data.get("unfair_advantage", []), LEAN_COLORS["unfair_advantage"])
        channels = create_block("Canaux", lean_data.get("channels", []), LEAN_COLORS["channels"])
        cust_segments = create_block("Segments Clients", lean_data.get("customer_segments", {}), LEAN_COLORS["customer_segments"])
        cost = create_block("Structure des Coûts", lean_data.get("cost_structure", {}), LEAN_COLORS["cost_structure"])
        revenue = create_block("Sources de Revenus", lean_data.get("revenue_streams", {}), LEAN_COLORS["revenue_streams"])
        
        lean_table = Table([
            # Row 1
            [problem, solution, uvp, unfair_adv, cust_segments],
            # Row 2
            [existing_alt, key_metrics, high_level, "", ""],
            # Row 3 - Channels (spans middle columns)
            ["", channels, "", "", ""],
            # Row 4 - Cost & Revenue
            [cost, "", "", revenue, ""],
        ], colWidths=[col_width] * 5)
        
        lean_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor(BIZGEN_DARK)),
            ('SPAN', (0, 0), (0, 1)),  # Problem spans 2 rows
            ('SPAN', (3, 0), (3, 1)),  # Unfair advantage spans 2 rows
            ('SPAN', (4, 0), (4, 1)),  # Customer segments spans 2 rows
            ('SPAN', (1, 2), (4, 2)),  # Channels spans 4 columns
            ('SPAN', (0, 3), (2, 3)),  # Cost spans 3 columns
            ('SPAN', (3, 3), (4, 3)),  # Revenue spans 2 columns
        ]))
        
        story.append(lean_table)
        
        # Footer
        story.append(Spacer(1, 0.5*cm))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
        story.append(Paragraph(f"Généré par BizGen AI - {datetime.now().strftime('%d/%m/%Y')}", footer_style))
        
        def on_page(canvas_obj, doc):
            self._add_header_footer(canvas_obj, doc, project_name, is_free_plan)
            if is_free_plan:
                self._add_watermark(canvas_obj, doc)
        
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_business_plan_pdf(
        self, 
        bp_data: Dict[str, Any], 
        project_name: str,
        is_free_plan: bool = False
    ) -> bytes:
        """Generate complete Business Plan as PDF with improved layout"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2.5*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles with professional appearance
        title_style = ParagraphStyle(
            'TitleCustom',
            parent=styles['Heading1'],
            fontSize=28,
            alignment=TA_CENTER,
            spaceAfter=30,
            textColor=colors.HexColor(BIZGEN_PRIMARY),
            fontName=FONT_TITLE
        )
        
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=14,
            alignment=TA_CENTER,
            textColor=colors.grey,
            spaceAfter=40
        )
        
        heading1_style = ParagraphStyle(
            'H1Custom',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor(BIZGEN_PRIMARY),
            spaceBefore=25,
            spaceAfter=12,
            fontName=FONT_HEADING,
            borderPadding=(0, 0, 5, 0),
        )
        
        heading2_style = ParagraphStyle(
            'H2Custom',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor(BIZGEN_SECONDARY),
            spaceBefore=15,
            spaceAfter=8,
            fontName=FONT_HEADING
        )
        
        body_style = ParagraphStyle(
            'BodyCustom',
            parent=styles['Normal'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceBefore=4,
            spaceAfter=4,
            leading=14,
            fontName=FONT_BODY
        )
        
        bullet_style = ParagraphStyle(
            'BulletCustom',
            parent=body_style,
            leftIndent=20,
            bulletIndent=10,
            spaceBefore=2,
            spaceAfter=2
        )
        
        story = []
        
        # ============ COVER PAGE ============
        story.append(Spacer(1, 4*cm))
        
        # Logo placeholder (could be an actual image)
        logo_style = ParagraphStyle('Logo', fontSize=48, alignment=TA_CENTER, textColor=colors.HexColor(BIZGEN_PRIMARY))
        story.append(Paragraph("📊", logo_style))
        story.append(Spacer(1, 1*cm))
        
        # Title
        story.append(Paragraph("BUSINESS PLAN", title_style))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(project_name, ParagraphStyle(
            'ProjectName',
            fontSize=22,
            alignment=TA_CENTER,
            textColor=colors.HexColor(BIZGEN_DARK),
            fontName=FONT_HEADING
        )))
        story.append(Spacer(1, 2*cm))
        
        # Metadata
        story.append(Paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}", subtitle_style))
        story.append(Paragraph("Généré par BizGen AI", subtitle_style))
        
        if is_free_plan:
            story.append(Spacer(1, 1*cm))
            warning_style = ParagraphStyle('Warning', fontSize=12, alignment=TA_CENTER, textColor=colors.HexColor('#EF4444'))
            story.append(Paragraph("⚠️ VERSION GRATUITE - Certaines fonctionnalités sont limitées", warning_style))
        
        story.append(PageBreak())
        
        # ============ TABLE OF CONTENTS ============
        story.append(Paragraph("Table des Matières", heading1_style))
        story.append(Spacer(1, 0.3*cm))
        
        toc_items = [
            ("1.", "Résumé Exécutif"),
            ("2.", "Présentation de l'Entreprise"),
            ("3.", "Analyse du Marché"),
            ("4.", "Analyse Concurrentielle"),
            ("5.", "Analyse SWOT"),
            ("6.", "Stratégie Marketing"),
            ("7.", "Plan Opérationnel"),
            ("8.", "Projections Financières"),
            ("9.", "Équipe"),
            ("10.", "Analyse des Risques"),
        ]
        
        toc_style = ParagraphStyle('TOC', parent=body_style, fontSize=12, spaceBefore=8)
        for num, item in toc_items:
            story.append(Paragraph(f"{num} {item}", toc_style))
        
        story.append(PageBreak())
        
        # ============ CONTENT SECTIONS ============
        
        # Helper function to add section
        def add_section(number: int, title: str, content_func):
            story.append(Paragraph(f"{number}. {title}", heading1_style))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor(BIZGEN_PRIMARY), spaceAfter=10))
            content_func()
            story.append(Spacer(1, 0.5*cm))
        
        # 1. Executive Summary
        def exec_summary_content():
            exec_summary = bp_data.get("executiveSummary", bp_data.get("executive_summary", "Non disponible"))
            story.append(Paragraph(exec_summary, body_style))
        
        add_section(1, "Résumé Exécutif", exec_summary_content)
        
        # 2. Company Overview
        def company_content():
            company = bp_data.get("companyOverview", bp_data.get("company_overview", {}))
            if company:
                if "mission" in company:
                    story.append(Paragraph("Mission", heading2_style))
                    story.append(Paragraph(company["mission"], body_style))
                if "vision" in company:
                    story.append(Paragraph("Vision", heading2_style))
                    story.append(Paragraph(company["vision"], body_style))
                if "values" in company:
                    story.append(Paragraph("Valeurs", heading2_style))
                    for value in company["values"]:
                        story.append(Paragraph(f"• {value}", bullet_style))
                if "description" in company:
                    story.append(Paragraph("Description", heading2_style))
                    story.append(Paragraph(company["description"], body_style))
        
        add_section(2, "Présentation de l'Entreprise", company_content)
        
        # 3. Market Analysis
        def market_content():
            market = bp_data.get("marketAnalysis", bp_data.get("market_analysis", {}))
            if market:
                field_mapping = [
                    ("industryOverview", "Aperçu du Secteur"),
                    ("targetMarket", "Marché Cible"),
                    ("marketSize", "Taille du Marché"),
                    ("trends", "Tendances du Marché"),
                    ("growth_potential", "Potentiel de Croissance"),
                ]
                for key, label in field_mapping:
                    if key in market:
                        story.append(Paragraph(label, heading2_style))
                        story.append(Paragraph(str(market[key]), body_style))
        
        add_section(3, "Analyse du Marché", market_content)
        
        # 4. Competitive Analysis
        def competitive_content():
            competitive = bp_data.get("competitiveAnalysis", bp_data.get("competitive_analysis", {}))
            if competitive:
                if "directCompetitors" in competitive:
                    story.append(Paragraph("Concurrents Directs", heading2_style))
                    for comp in competitive["directCompetitors"]:
                        story.append(Paragraph(f"• {comp}", bullet_style))
                if "indirectCompetitors" in competitive:
                    story.append(Paragraph("Concurrents Indirects", heading2_style))
                    for comp in competitive["indirectCompetitors"]:
                        story.append(Paragraph(f"• {comp}", bullet_style))
                if "competitiveAdvantage" in competitive:
                    story.append(Paragraph("Avantage Concurrentiel", heading2_style))
                    story.append(Paragraph(competitive["competitiveAdvantage"], body_style))
        
        add_section(4, "Analyse Concurrentielle", competitive_content)
        
        # 5. SWOT Analysis
        def swot_content():
            swot = bp_data.get("swot", {})
            if swot:
                swot_data = [
                    ["Forces", "Faiblesses"],
                    ["\n".join(["• " + s for s in swot.get("strengths", [])]), 
                     "\n".join(["• " + w for w in swot.get("weaknesses", [])])],
                    ["Opportunités", "Menaces"],
                    ["\n".join(["• " + o for o in swot.get("opportunities", [])]),
                     "\n".join(["• " + t for t in swot.get("threats", [])])]
                ]
                
                swot_table = Table(swot_data, colWidths=[8*cm, 8*cm])
                swot_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BIZGEN_PRIMARY)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), FONT_HEADING),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('TOPPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 2), (-1, 2), colors.HexColor(BIZGEN_SECONDARY)),
                    ('TEXTCOLOR', (0, 2), (-1, 2), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor(BIZGEN_DARK)),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('TOPPADDING', (0, 1), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ]))
                story.append(swot_table)
        
        add_section(5, "Analyse SWOT", swot_content)
        
        # 6. Marketing Strategy
        def marketing_content():
            marketing = bp_data.get("marketingStrategy", bp_data.get("marketing_strategy", {}))
            if marketing:
                field_mapping = [
                    ("positioning", "Positionnement"),
                    ("pricing", "Stratégie de Prix"),
                    ("distribution", "Distribution"),
                ]
                for key, label in field_mapping:
                    if key in marketing:
                        story.append(Paragraph(label, heading2_style))
                        story.append(Paragraph(str(marketing[key]), body_style))
                
                if "channels" in marketing:
                    story.append(Paragraph("Canaux de Distribution", heading2_style))
                    for channel in marketing["channels"]:
                        story.append(Paragraph(f"• {channel}", bullet_style))
        
        add_section(6, "Stratégie Marketing", marketing_content)
        
        # 7. Operations Plan
        def operations_content():
            operations = bp_data.get("operationsPlan", bp_data.get("operations_plan", {}))
            if operations:
                if "keyActivities" in operations:
                    story.append(Paragraph("Activités Clés", heading2_style))
                    for activity in operations["keyActivities"]:
                        story.append(Paragraph(f"• {activity}", bullet_style))
                if "milestones" in operations:
                    story.append(Paragraph("Jalons", heading2_style))
                    milestone_data = [["Jalon", "Date"]]
                    for m in operations["milestones"]:
                        milestone_data.append([str(m.get("name", m)), str(m.get("date", ""))])
                    
                    if len(milestone_data) > 1:
                        m_table = Table(milestone_data, colWidths=[12*cm, 4*cm])
                        m_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BIZGEN_PRIMARY)),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                            ('FONTNAME', (0, 0), (-1, 0), FONT_HEADING),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ]))
                        story.append(m_table)
        
        add_section(7, "Plan Opérationnel", operations_content)
        
        # 8. Financial Projections
        def financial_content():
            financial = bp_data.get("financialProjections", bp_data.get("financial_projections", {}))
            if financial:
                financial_data = [
                    ["Indicateur", "Valeur"],
                    ["Revenu Année 1", financial.get("year1Revenue", "N/A")],
                    ["Revenu Année 2", financial.get("year2Revenue", "N/A")],
                    ["Revenu Année 3", financial.get("year3Revenue", "N/A")],
                    ["Seuil de Rentabilité", f"Mois {financial.get('breakEvenMonth', 'N/A')}"],
                    ["Financement Requis", financial.get("fundingRequired", "N/A")],
                ]
                
                fin_table = Table(financial_data, colWidths=[8*cm, 8*cm])
                fin_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(BIZGEN_PRIMARY)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), FONT_HEADING),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor(BIZGEN_LIGHT)]),
                ]))
                story.append(fin_table)
        
        add_section(8, "Projections Financières", financial_content)
        
        # 9. Team
        def team_content():
            team = bp_data.get("team", {})
            if team:
                if "founders" in team:
                    story.append(Paragraph("Fondateurs", heading2_style))
                    for founder in team["founders"]:
                        story.append(Paragraph(f"• {founder}", bullet_style))
                if "keyMembers" in team:
                    story.append(Paragraph("Membres Clés", heading2_style))
                    for member in team["keyMembers"]:
                        story.append(Paragraph(f"• {member}", bullet_style))
        
        add_section(9, "Équipe", team_content)
        
        # 10. Risk Analysis
        def risk_content():
            risk = bp_data.get("riskAnalysis", bp_data.get("risk_analysis", {}))
            if risk:
                if "risks" in risk:
                    story.append(Paragraph("Risques Identifiés", heading2_style))
                    for r in risk["risks"]:
                        story.append(Paragraph(f"• {r}", bullet_style))
                if "mitigations" in risk:
                    story.append(Paragraph("Stratégies d'Atténuation", heading2_style))
                    for m in risk["mitigations"]:
                        story.append(Paragraph(f"• {m}", bullet_style))
        
        add_section(10, "Analyse des Risques", risk_content)
        
        # Footer note
        story.append(Spacer(1, 1*cm))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor(BIZGEN_PRIMARY)))
        story.append(Spacer(1, 0.3*cm))
        footer_text = f"Document généré par BizGen AI le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        story.append(Paragraph(footer_text, ParagraphStyle('FooterNote', fontSize=9, textColor=colors.grey, alignment=TA_CENTER)))
        
        # Build PDF
        def on_page(canvas_obj, doc):
            self._add_header_footer(canvas_obj, doc, project_name, is_free_plan)
            if is_free_plan:
                self._add_watermark(canvas_obj, doc)
        
        doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
        buffer.seek(0)
        return buffer.getvalue()

    # ============================================
    # DOCX GENERATION - IMPROVED
    # ============================================
    
    def _set_cell_shading(self, cell, color: str):
        """Set background color for a table cell"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color.replace('#', ''))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_page_break(self, doc: Document):
        """Add a page break"""
        doc.add_page_break()
    
    def _create_toc(self, doc: Document):
        """Add a table of contents field"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        run = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)
        
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
    
    def _set_document_styles(self, doc: Document):
        """Set up professional document styles"""
        styles = doc.styles
        
        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        normal.paragraph_format.space_after = Pt(8)
        
        # Heading 1
        h1 = styles['Heading 1']
        h1.font.name = 'Calibri Light'
        h1.font.size = Pt(24)
        h1.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # BIZGEN_PRIMARY
        h1.font.bold = True
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(12)
        
        # Heading 2
        h2 = styles['Heading 2']
        h2.font.name = 'Calibri Light'
        h2.font.size = Pt(16)
        h2.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)  # BIZGEN_SECONDARY
        h2.font.bold = True
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)
        
        # Heading 3
        h3 = styles['Heading 3']
        h3.font.name = 'Calibri'
        h3.font.size = Pt(14)
        h3.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)  # BIZGEN_DARK
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
    
    def _add_header_footer_docx(self, doc: Document, project_name: str, is_free_plan: bool = False):
        """Add header and footer to all sections"""
        for section in doc.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = header_para.add_run(f"BizGen AI - {project_name}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            run.font.bold = True
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = footer_para.add_run(f"Document généré par BizGen AI le {datetime.now().strftime('%d/%m/%Y')}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            
            if is_free_plan:
                footer_para.add_run(" | ")
                run = footer_para.add_run("VERSION GRATUITE")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                run.font.bold = True
    
    def generate_business_plan_docx(
        self, 
        bp_data: Dict[str, Any], 
        project_name: str,
        is_free_plan: bool = False
    ) -> bytes:
        """Generate Business Plan as Word document with professional styling"""
        doc = Document()
        
        # Set document styles
        self._set_document_styles(doc)
        
        # Set document margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # ============ COVER PAGE ============
        doc.add_paragraph()  # Spacing
        doc.add_paragraph()
        
        # Logo/Icon
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run("📊")
        logo_run.font.size = Pt(48)
        
        doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("BUSINESS PLAN")
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        
        doc.add_paragraph()
        
        # Project name
        project_para = doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_run = project_para.add_run(project_name)
        project_run.font.size = Pt(24)
        project_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Date: {datetime.now().strftime('%d %B %Y')}")
        date_run.font.size = Pt(14)
        date_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        
        # Generated by
        gen_para = doc.add_paragraph()
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_run = gen_para.add_run("Généré par BizGen AI")
        gen_run.font.size = Pt(12)
        gen_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        
        if is_free_plan:
            doc.add_paragraph()
            warning_para = doc.add_paragraph()
            warning_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            warning_run = warning_para.add_run("⚠️ VERSION GRATUITE")
            warning_run.font.size = Pt(12)
            warning_run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
            warning_run.font.bold = True
        
        self._add_page_break(doc)
        
        # ============ TABLE OF CONTENTS ============
        doc.add_heading("Table des Matières", level=1)
        
        toc_items = [
            "1. Résumé Exécutif",
            "2. Présentation de l'Entreprise",
            "3. Analyse du Marché",
            "4. Analyse Concurrentielle",
            "5. Analyse SWOT",
            "6. Stratégie Marketing",
            "7. Plan Opérationnel",
            "8. Projections Financières",
            "9. Équipe",
            "10. Analyse des Risques",
        ]
        
        for item in toc_items:
            toc_para = doc.add_paragraph(item)
            toc_para.paragraph_format.left_indent = Cm(1)
        
        self._add_page_break(doc)
        
        # ============ CONTENT ============
        
        # 1. Executive Summary
        doc.add_heading("1. Résumé Exécutif", level=1)
        exec_summary = bp_data.get("executiveSummary", bp_data.get("executive_summary", "Non disponible"))
        doc.add_paragraph(exec_summary)
        
        # 2. Company Overview
        company = bp_data.get("companyOverview", bp_data.get("company_overview", {}))
        if company:
            doc.add_heading("2. Présentation de l'Entreprise", level=1)
            
            if "mission" in company:
                doc.add_heading("Mission", level=2)
                doc.add_paragraph(company["mission"])
            
            if "vision" in company:
                doc.add_heading("Vision", level=2)
                doc.add_paragraph(company["vision"])
            
            if "values" in company:
                doc.add_heading("Valeurs", level=2)
                for value in company["values"]:
                    para = doc.add_paragraph(value, style='List Bullet')
            
            if "description" in company:
                doc.add_heading("Description", level=2)
                doc.add_paragraph(company["description"])
        
        # 3. Market Analysis
        market = bp_data.get("marketAnalysis", bp_data.get("market_analysis", {}))
        if market:
            doc.add_heading("3. Analyse du Marché", level=1)
            
            field_mapping = [
                ("industryOverview", "Aperçu du Secteur"),
                ("targetMarket", "Marché Cible"),
                ("marketSize", "Taille du Marché"),
                ("trends", "Tendances"),
            ]
            
            for key, label in field_mapping:
                if key in market:
                    doc.add_heading(label, level=2)
                    doc.add_paragraph(str(market[key]))
        
        # 4. Competitive Analysis
        competitive = bp_data.get("competitiveAnalysis", bp_data.get("competitive_analysis", {}))
        if competitive:
            doc.add_heading("4. Analyse Concurrentielle", level=1)
            
            if "directCompetitors" in competitive:
                doc.add_heading("Concurrents Directs", level=2)
                for comp in competitive["directCompetitors"]:
                    doc.add_paragraph(comp, style='List Bullet')
            
            if "competitiveAdvantage" in competitive:
                doc.add_heading("Avantage Concurrentiel", level=2)
                doc.add_paragraph(competitive["competitiveAdvantage"])
        
        # 5. SWOT Analysis
        swot = bp_data.get("swot", {})
        if swot:
            doc.add_heading("5. Analyse SWOT", level=1)
            
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Set column widths
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Cm(8)
            
            # Headers
            cells = table.rows[0].cells
            cells[0].text = "Forces"
            cells[1].text = "Faiblesses"
            self._set_cell_shading(cells[0], BIZGEN_PRIMARY)
            self._set_cell_shading(cells[1], BIZGEN_PRIMARY)
            for cell in cells:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Strengths & Weaknesses
            cells = table.rows[1].cells
            cells[0].text = "\n".join(["• " + s for s in swot.get("strengths", [])])
            cells[1].text = "\n".join(["• " + w for w in swot.get("weaknesses", [])])
            
            # Opportunities & Threats headers
            cells = table.rows[2].cells
            cells[0].text = "Opportunités"
            cells[1].text = "Menaces"
            self._set_cell_shading(cells[0], BIZGEN_SECONDARY)
            self._set_cell_shading(cells[1], BIZGEN_SECONDARY)
            for cell in cells:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Opportunities & Threats content
            cells = table.rows[3].cells
            cells[0].text = "\n".join(["• " + o for o in swot.get("opportunities", [])])
            cells[1].text = "\n".join(["• " + t for t in swot.get("threats", [])])
        
        # 6. Marketing Strategy
        marketing = bp_data.get("marketingStrategy", bp_data.get("marketing_strategy", {}))
        if marketing:
            doc.add_heading("6. Stratégie Marketing", level=1)
            
            if "positioning" in marketing:
                doc.add_heading("Positionnement", level=2)
                doc.add_paragraph(marketing["positioning"])
            
            if "channels" in marketing:
                doc.add_heading("Canaux de Distribution", level=2)
                for channel in marketing["channels"]:
                    doc.add_paragraph(channel, style='List Bullet')
        
        # 7. Operations Plan
        operations = bp_data.get("operationsPlan", bp_data.get("operations_plan", {}))
        if operations:
            doc.add_heading("7. Plan Opérationnel", level=1)
            
            if "keyActivities" in operations:
                doc.add_heading("Activités Clés", level=2)
                for activity in operations["keyActivities"]:
                    doc.add_paragraph(activity, style='List Bullet')
            
            if "milestones" in operations:
                doc.add_heading("Jalons", level=2)
                for milestone in operations["milestones"]:
                    doc.add_paragraph(str(milestone), style='List Bullet')
        
        # 8. Financial Projections
        financial = bp_data.get("financialProjections", bp_data.get("financial_projections", {}))
        if financial:
            doc.add_heading("8. Projections Financières", level=1)
            
            table = doc.add_table(rows=6, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            data = [
                ("Indicateur", "Valeur"),
                ("Revenu Année 1", str(financial.get("year1Revenue", "N/A"))),
                ("Revenu Année 2", str(financial.get("year2Revenue", "N/A"))),
                ("Revenu Année 3", str(financial.get("year3Revenue", "N/A"))),
                ("Seuil de Rentabilité", f"Mois {financial.get('breakEvenMonth', 'N/A')}"),
                ("Financement Requis", str(financial.get("fundingRequired", "N/A"))),
            ]
            
            for i, (key, value) in enumerate(data):
                row = table.rows[i]
                row.cells[0].text = key
                row.cells[1].text = value
                
                # Style header row
                if i == 0:
                    self._set_cell_shading(row.cells[0], BIZGEN_PRIMARY)
                    self._set_cell_shading(row.cells[1], BIZGEN_PRIMARY)
                    for cell in row.cells:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        cell.paragraphs[0].runs[0].font.bold = True
                
                # Alternate row colors
                elif i % 2 == 0:
                    self._set_cell_shading(row.cells[0], BIZGEN_LIGHT)
                    self._set_cell_shading(row.cells[1], BIZGEN_LIGHT)
        
        # 9. Team
        team = bp_data.get("team", {})
        if team:
            doc.add_heading("9. Équipe", level=1)
            
            if "founders" in team:
                doc.add_heading("Fondateurs", level=2)
                for founder in team["founders"]:
                    doc.add_paragraph(founder, style='List Bullet')
        
        # 10. Risk Analysis
        risk = bp_data.get("riskAnalysis", bp_data.get("risk_analysis", {}))
        if risk:
            doc.add_heading("10. Analyse des Risques", level=1)
            
            if "risks" in risk:
                doc.add_heading("Risques Identifiés", level=2)
                for r in risk["risks"]:
                    doc.add_paragraph(r, style='List Bullet')
            
            if "mitigations" in risk:
                doc.add_heading("Stratégies d'Atténuation", level=2)
                for m in risk["mitigations"]:
                    doc.add_paragraph(m, style='List Bullet')
        
        # Add header and footer
        self._add_header_footer_docx(doc, project_name, is_free_plan)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    # ============================================
    # ZIP EXPORT
    # ============================================
    
    def generate_export_zip(
        self,
        project_name: str,
        bmc_data: Optional[Dict[str, Any]] = None,
        lean_data: Optional[Dict[str, Any]] = None,
        bp_data: Optional[Dict[str, Any]] = None,
        include_png: bool = True,
        include_pdf: bool = True,
        include_docx: bool = True,
        is_free_plan: bool = False
    ) -> bytes:
        """
        Generate a ZIP file containing all export documents
        
        Args:
            project_name: Name of the project
            bmc_data: Business Model Canvas data
            lean_data: Lean Canvas data
            bp_data: Business Plan data
            include_png: Include PNG exports
            include_pdf: Include PDF exports
            include_docx: Include DOCX exports
            is_free_plan: Whether the user is on a free plan
        
        Returns:
            ZIP file as bytes
        """
        buffer = io.BytesIO()
        
        # Sanitize project name for filenames
        safe_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in project_name)
        safe_name = safe_name[:50]  # Limit length
        
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # BMC exports
            if bmc_data:
                if include_png:
                    png_data = self.generate_bmc_png(bmc_data, project_name)
                    zipf.writestr(f"{safe_name}_BMC.png", png_data)
                
                if include_pdf:
                    pdf_data = self.generate_bmc_pdf(bmc_data, project_name, is_free_plan)
                    zipf.writestr(f"{safe_name}_BMC.pdf", pdf_data)
            
            # Lean Canvas exports
            if lean_data:
                if include_png:
                    png_data = self.generate_lean_canvas_png(lean_data, project_name)
                    zipf.writestr(f"{safe_name}_Lean_Canvas.png", png_data)
                
                if include_pdf:
                    pdf_data = self.generate_lean_canvas_pdf(lean_data, project_name, is_free_plan)
                    zipf.writestr(f"{safe_name}_Lean_Canvas.pdf", pdf_data)
            
            # Business Plan exports
            if bp_data:
                if include_pdf:
                    pdf_data = self.generate_business_plan_pdf(bp_data, project_name, is_free_plan)
                    zipf.writestr(f"{safe_name}_Business_Plan.pdf", pdf_data)
                
                if include_docx:
                    docx_data = self.generate_business_plan_docx(bp_data, project_name, is_free_plan)
                    zipf.writestr(f"{safe_name}_Business_Plan.docx", docx_data)
            
            # Add README
            readme_content = f"""BizGen AI Export Package
========================

Project: {project_name}
Generated: {datetime.now().strftime('%d/%m/%Y à %H:%M')}

Contents:
---------
"""
            if bmc_data:
                readme_content += f"- {safe_name}_BMC.png (Business Model Canvas - Image)\n"
                readme_content += f"- {safe_name}_BMC.pdf (Business Model Canvas - PDF)\n"
            if lean_data:
                readme_content += f"- {safe_name}_Lean_Canvas.png (Lean Canvas - Image)\n"
                readme_content += f"- {safe_name}_Lean_Canvas.pdf (Lean Canvas - PDF)\n"
            if bp_data:
                readme_content += f"- {safe_name}_Business_Plan.pdf (Business Plan - PDF)\n"
                readme_content += f"- {safe_name}_Business_Plan.docx (Business Plan - Word)\n"
            
            if is_free_plan:
                readme_content += "\n--- FREE VERSION ---\n"
                readme_content += "Upgrade to Pro for unlimited exports and additional features!\n"
            
            readme_content += "\nGenerated by BizGen AI - https://bizgen.ai\n"
            zipf.writestr("README.txt", readme_content)
        
        buffer.seek(0)
        return buffer.getvalue()

    # ============================================
    # MAIN EXPORT METHOD
    # ============================================
    
    def export_document(
        self,
        doc_type: str,  # bmc, lean, bp
        format_type: str,  # pdf, docx, png
        data: Dict[str, Any],
        project_name: str,
        is_free_plan: bool = False
    ) -> bytes:
        """Main export method"""
        
        if doc_type == "bmc":
            if format_type == "pdf":
                return self.generate_bmc_pdf(data, project_name, is_free_plan)
            elif format_type == "png":
                return self.generate_bmc_png(data, project_name)
            else:
                raise ValueError(f"Format {format_type} not supported for BMC")
        
        elif doc_type == "lean":
            if format_type == "pdf":
                return self.generate_lean_canvas_pdf(data, project_name, is_free_plan)
            elif format_type == "png":
                return self.generate_lean_canvas_png(data, project_name)
            else:
                raise ValueError(f"Format {format_type} not supported for Lean Canvas")
        
        elif doc_type == "bp":
            if format_type == "pdf":
                return self.generate_business_plan_pdf(data, project_name, is_free_plan)
            elif format_type == "docx":
                return self.generate_business_plan_docx(data, project_name, is_free_plan)
            else:
                raise ValueError(f"Format {format_type} not supported for Business Plan")
        
        else:
            raise ValueError(f"Document type {doc_type} not supported")
    
    def export_all(
        self,
        project_name: str,
        bmc_data: Optional[Dict[str, Any]] = None,
        lean_data: Optional[Dict[str, Any]] = None,
        bp_data: Optional[Dict[str, Any]] = None,
        is_free_plan: bool = False
    ) -> bytes:
        """
        Export all documents as a ZIP file
        
        Convenience method for exporting all canvas types and business plan
        """
        return self.generate_export_zip(
            project_name=project_name,
            bmc_data=bmc_data,
            lean_data=lean_data,
            bp_data=bp_data,
            include_png=True,
            include_pdf=True,
            include_docx=True,
            is_free_plan=is_free_plan
        )


# Singleton instance
export_service = ExportService()
